from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
SOURCE_MD = BASE_DIR / "2026-06-02-meltblown-vision-redesign-plan.md"
OUT_DOCX = BASE_DIR / "meltblown_vision_redesign_plan.docx"

FONT_LATIN = "Calibri"
FONT_EAST_ASIA = "Microsoft YaHei"
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x0B, 0x25, 0x45)
MUTED = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "F2F4F7"
BORDER = "B7C9DD"


def set_run_font(run, size=None, bold=None, color=None):
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_paragraph_font(paragraph, size=11, color=None):
    for run in paragraph.runs:
        set_run_font(run, size=size, color=color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")

    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = tbl_borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tbl_borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "6")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def table_widths(headers):
    joined = "|".join(headers)
    count = len(headers)
    if count == 2:
        return [2600, 6760]
    if count == 3:
        if "验收时间" in joined:
            return [1800, 2700, 4860]
        if "风险项" in joined:
            return [2200, 5000, 2160]
        if "角色" in joined:
            return [1850, 3450, 4060]
        return [2100, 4260, 3000]
    if count == 4:
        if "软件设计与开发主线" in joined:
            return [1400, 3200, 3000, 1760]
        if "可并行工作" in joined:
            return [2300, 1500, 1900, 3660]
        if "模块" in joined:
            return [1900, 3500, 3060, 900]
        return [1400, 2200, 4000, 1760]
    base = 9360 // count
    widths = [base] * count
    widths[-1] += 9360 - sum(widths)
    return widths


def add_runs_with_bold(paragraph, text, size=11, color=None):
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        clean = part[2:-2] if is_bold else part
        clean = clean.replace("`", "")
        run = paragraph.add_run(clean)
        set_run_font(run, size=size, bold=is_bold, color=color)


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        rows.append(row)
        i += 1
    if len(rows) >= 2 and all(set(c.replace(" ", "")) <= {"-"} for c in rows[1]):
        rows.pop(1)
    return rows, i


def add_table(doc, rows):
    if not rows:
        return
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    widths = table_widths(rows[0])
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.08
            add_runs_with_bold(paragraph, value, size=8.4 if len(rows[0]) >= 4 else 9)
            if r_idx == 0:
                set_cell_shading(cell, HEADER_FILL)
                for run in paragraph.runs:
                    set_run_font(run, size=9, bold=True, color=INK)
            elif len(value) <= 8 or value in {"必须", "应完成", "高", "中", "低", "无"}:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_table_width(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def setup_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.start_type = WD_SECTION_START.NEW_PAGE

    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_header(doc):
    header = doc.sections[0].header
    paragraph = header.paragraphs[0]
    paragraph.text = ""
    run = paragraph.add_run("熔喷视觉检测软件重新设计计划")
    set_run_font(run, size=9, color=MUTED)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def build_doc():
    text = SOURCE_MD.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    setup_styles(doc)
    add_header(doc)

    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line or line == "---":
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        if line.startswith("# "):
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_after = Pt(8)
            run = paragraph.add_run(line[2:].strip())
            set_run_font(run, size=22, bold=True, color=INK)
            i += 1
            continue
        if line.startswith("## "):
            paragraph = doc.add_paragraph(line[3:].strip(), style="Heading 1")
            set_paragraph_font(paragraph, size=16, color=BLUE)
            i += 1
            continue
        if line.startswith("### "):
            paragraph = doc.add_paragraph(line[4:].strip(), style="Heading 2")
            set_paragraph_font(paragraph, size=13, color=BLUE)
            i += 1
            continue
        if line.startswith("> "):
            paragraph = doc.add_paragraph()
            paragraph.paragraph_format.left_indent = Inches(0.18)
            paragraph.paragraph_format.space_after = Pt(8)
            add_runs_with_bold(paragraph, line[2:].strip(), size=10, color=MUTED)
            i += 1
            continue
        if line.startswith("- "):
            paragraph = doc.add_paragraph(style="List Bullet")
            paragraph.paragraph_format.space_after = Pt(4)
            paragraph.paragraph_format.line_spacing = 1.167
            add_runs_with_bold(paragraph, line[2:].strip(), size=10.5)
            i += 1
            continue

        paragraph = doc.add_paragraph()
        add_runs_with_bold(paragraph, line, size=11)
        i += 1

    doc.core_properties.title = "熔喷视觉检测软件重新设计计划"
    doc.core_properties.subject = "软件设计计划与并行工作分析"
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
