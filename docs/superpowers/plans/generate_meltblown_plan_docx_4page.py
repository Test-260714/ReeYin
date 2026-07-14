from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
OUT_DOCX = BASE_DIR / "熔喷视觉检测软件重新设计计划_4页版.docx"

FONT_LATIN = "Calibri"
FONT_EAST_ASIA = "Microsoft YaHei"
INK = RGBColor(0x0B, 0x25, 0x45)
BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "E8EEF5"
LIGHT_FILL = "F4F6F9"
BORDER = "B7C9DD"


def set_run_font(run, size=10, bold=False, color=None):
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    run.font.size = Pt(size)
    run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=60, bottom=60, start=90, end=90):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), BORDER)


def set_table_geometry(table, widths_in):
    widths = [int(width * 1440) for width in widths_in]
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths_in[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
    set_table_borders(table)


def add_paragraph(doc, text="", size=10.2, bold=False, color=None, align=None, space_after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.08
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6 if level == 1 else 3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    set_run_font(run, size=13 if level == 1 else 11, bold=True, color=BLUE if level == 1 else DARK_BLUE)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = 1.05
        run = p.add_run(item)
        set_run_font(run, size=9.4)


def add_table(doc, headers, rows, widths_in, font_size=8.2, header_size=8.4):
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = ""
        p = hdr[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(header)
        set_run_font(run, size=header_size, bold=True, color=INK)
        set_cell_shading(hdr[idx], HEADER_FILL)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = ""
            p = cells[idx].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            if value in {"必须", "应完成", "高", "中", "低"}:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(value)
            set_run_font(run, size=font_size)
    set_table_geometry(table, widths_in)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def setup_doc():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(3)
    normal.paragraph_format.line_spacing = 1.08

    header = section.header.paragraphs[0]
    header.text = ""
    run = header.add_run("熔喷视觉检测软件重新设计计划")
    set_run_font(run, size=8.5, color=MUTED)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    return doc


def page_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(title)
    set_run_font(run, size=19, bold=True, color=INK)
    if subtitle:
        add_paragraph(doc, subtitle, size=10.5, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=5)


def build_doc():
    doc = setup_doc()

    # Page 1: objectives and scope
    page_title(doc, "熔喷视觉检测软件重新设计计划", "核心功能整改版 / 现场验证版")
    add_table(
        doc,
        ["项目项", "计划说明"],
        [
            ["项目对象", "熔喷视觉检测设备配套软件"],
            ["周期定位", "约两周、10 个工作日；本版不列具体起止日期"],
            ["交付定位", "核心功能整改版、现场验证版，后续根据现场连续运行继续优化"],
            ["设计原则", "复用现有相机 SDK、通讯接口、检测算法接口和关键配置，优先解决现场急需问题"],
        ],
        [1.35, 5.85],
        font_size=8.6,
    )
    add_heading(doc, "一、设计目标", 1)
    add_paragraph(
        doc,
        "本次重新设计不从零开发完整视觉平台，而是在现有软件和现场流程基础上重设检测业务模块，重点形成可运行、可追溯、可验证的现场版本。",
        size=9.8,
    )
    add_table(
        doc,
        ["目标方向", "关键设计内容", "交付 / 验收要点"],
        [
            ["状态与报警", "相机在线、离线、异常、断连状态统一建模；异常进入报警闭环", "断连可发现、可报警、可记录、可恢复"],
            ["数据与批号", "检测数据、异常数据、批号数据统一保存；批号唯一性和修改审计", "不覆盖、不误删、有权限控制和操作记录"],
            ["幅宽追溯", "幅宽上下限判断、异常报警、本地持久化、按批号和时间查询", "幅宽异常可报警、可追溯、可复盘"],
            ["报表统计", "按时间范围、批号、产品和异常类型查询汇总并导出", "满足生产复盘和质量追溯"],
            ["操作优化", "高频参数前置、参数分组、模板化配置、常见操作简化", "降低误操作，提高现场使用效率"],
            ["算法验证", "多相机标定、AI 推理、缺陷位置计算和结果输出串联", "算法链路现场可运行、结果可验证"],
        ],
        [1.28, 3.35, 2.57],
        font_size=8.0,
    )
    add_heading(doc, "二、范围边界", 1)
    add_bullets(
        doc,
        [
            "优先完成现场急需功能：相机状态、报警闭环、批号安全、幅宽追溯、报表统计、参数操作优化。",
            "不承诺在本周期内完成所有极端场景、复杂统计图表、多语言、多工位扩展和完整自动化测试平台。",
            "如旧软件无源码、相机链路不稳定、历史数据迁移复杂或检测算法需重写，应单独评估缓冲周期。",
        ],
    )

    doc.add_page_break()

    # Page 2: phased plan and parallel work
    page_title(doc, "阶段计划与并行安排")
    add_heading(doc, "三、阶段计划", 1)
    add_table(
        doc,
        ["阶段", "软件设计与开发主线", "算法 / 现场准备主线", "阶段输出"],
        [
            ["阶段一：启动与设计", "梳理旧软件功能、相机接口、数据库结构、生产流程和整改优先级；设计模块结构、状态模型、报警模型、批号模型、幅宽数据表和报表结构", "确认标定板、显卡、驱动、CUDA、推理环境；启动现场数据标注；明确算法输入输出格式", "需求清单、功能边界、风险清单、技术方案、数据库 / 配置方案"],
            ["阶段二：采集与报警闭环", "实现相机连接、心跳、在线 / 离线状态、断线报警、自动重连和日志；接入采图失败、黑屏 / 花屏、模糊、丢帧等异常判断", "持续数据标注和模型训练；同步验证推理环境和模型版本", "相机状态监控可用，图像异常报警闭环可验证"],
            ["阶段三：数据安全与追溯", "实现批号唯一性校验、冲突防护、受控修改、数据不覆盖、操作日志；实现幅宽上下限判断、异常报警、历史查询和追溯", "固化可部署模型版本，准备标定和缺陷定位现场验证数据", "批号冲突不导致数据丢失，幅宽数据可报警、可查询、可追溯"],
            ["阶段四：报表、界面与联调", "完成多维查询、异常统计、基础导出和查询性能优化；前置高频参数、参数分组和模板；开展全流程联调和关键问题修复", "部署训练模型和推理环境；完成多相机标定、AI 推理、缺陷位置计算和算法模块串联验证", "操作界面整改版、统计报表、初版说明文档、现场验证版"],
        ],
        [1.02, 2.68, 1.95, 1.55],
        font_size=7.35,
        header_size=7.8,
    )
    add_heading(doc, "四、可并行内容", 1)
    add_table(
        doc,
        ["并行工作", "前置条件", "并行说明"],
        [
            ["需求梳理、标定板准备、显卡升级确认", "无", "可由项目、算法、现场人员分别推进"],
            ["现场数据标注与软件整改", "现场数据可获取", "标注和训练从启动阶段开始，不等待软件全部完成"],
            ["架构设计与算法接口约定", "需求边界基本确认", "同步确定输入、输出和落库格式，减少联调返工"],
            ["相机状态监控与数据库基础结构", "数据模型完成", "采集链路和持久化框架可分人实现"],
            ["异常图像报警与批号数据安全", "报警模型和批号模型完成", "业务相对独立，但统一接入日志和审计"],
            ["幅宽追溯与报表查询", "幅宽数据表和查询维度完成", "保存接口先定后，报表统计可提前开发"],
            ["UI 操作优化与后台功能开发", "参数分组和操作清单完成", "页面结构先行，随后绑定后台接口"],
            ["说明文档与模块测试", "核心流程基本确定", "边开发边记录，模块完成后立即测试"],
        ],
        [2.35, 1.82, 3.03],
        font_size=7.85,
    )

    doc.add_page_break()

    # Page 3: dependencies, acceptance, risks
    page_title(doc, "依赖关系、验收与风险")
    add_heading(doc, "五、必须串行或强依赖内容", 1)
    add_table(
        doc,
        ["串行链路", "原因"],
        [
            ["需求确认 -> 架构与数据设计 -> 模块开发", "需求边界和数据结构未定时直接开发，容易造成返工"],
            ["相机 SDK 接入 -> 状态监控 -> 断连报警 -> 异常恢复测试", "报警和恢复依赖稳定的相机连接状态判断"],
            ["报警模型 -> 图像异常 / 幅宽异常 / 相机断连报警", "需要统一报警记录、确认、拦截或停机流程"],
            ["批号模型 -> 数据保存 -> 历史查询 -> 报表统计", "报表依赖准确、不可覆盖、可追溯的数据来源"],
            ["数据标注 -> 模型训练 -> 模型部署 -> AI 推理验证", "算法现场验证必须建立在可用模型和环境之上"],
            ["单模块验证 -> 全流程联调 -> 交付验收", "最终验收必须覆盖采集、检测、报警、存储、查询、导出全过程"],
        ],
        [3.35, 3.85],
        font_size=8.0,
    )
    add_heading(doc, "六、验收安排", 1)
    add_table(
        doc,
        ["验收节点", "验收重点", "通过标准"],
        [
            ["阶段验收一", "相机状态、断连报警、图像异常报警", "相机状态清晰可见，断连和图像异常可报警并记录"],
            ["阶段验收二", "批号防冲突、幅宽数据存储、报表查询", "批号不覆盖数据，幅宽可追溯，报表可按维度查询"],
            ["交付验收", "真实生产流程全链路", "采集、检测、报警、存储、查询、导出全过程跑通"],
            ["运行观察", "现场连续运行稳定性", "记录断连次数、报警准确性、查询性能和数据完整性"],
        ],
        [1.35, 2.35, 3.5],
        font_size=8.0,
    )
    add_heading(doc, "七、风险与缓冲", 1)
    add_table(
        doc,
        ["风险项", "影响说明", "处理建议"],
        [
            ["旧软件无源码或逻辑不清", "需额外梳理业务流程、数据库字段和接口调用", "增加梳理和验证缓冲"],
            ["相机 SDK 或硬件链路不稳定", "断连可能涉及供电、网线、网卡、触发信号和驱动", "软件与硬件联合排查"],
            ["历史数据迁移复杂", "旧检测记录、批号记录、幅宽数据需要单独校验", "迁移任务独立评估"],
            ["现场连续运行验证不足", "现场验证版不能替代长周期量产稳定性验证", "交付后持续观察并优化"],
            ["算法或检测规则重写", "原算法不可复用会显著拉长周期", "重新评估算法周期和人员投入"],
        ],
        [1.8, 3.35, 2.05],
        font_size=7.8,
    )

    doc.add_page_break()

    # Page 4: personnel assignment table
    page_title(doc, "人员工作分配表")
    add_paragraph(
        doc,
        "说明：以下表格用于指定对应人员负责相关工作。负责人和协作人员可在项目启动会后填写，交付标准用于后续跟踪和验收。",
        size=9.6,
        color=MUTED,
        space_after=5,
    )
    add_table(
        doc,
        ["工作模块", "主要工作", "负责人", "协作人员", "交付 / 验收标准"],
        [
            ["项目 / 需求管理", "需求边界、整改优先级、验收标准、对外沟通", "____", "____", "需求清单、风险清单、验收口径确认"],
            ["架构与数据设计", "模块边界、数据模型、配置方案、接口约定", "____", "____", "技术方案、数据库 / 配置方案确认"],
            ["相机采集与状态", "相机连接、心跳、在线 / 离线、断线重连、日志", "____", "____", "相机状态监控可用，断连可恢复"],
            ["报警闭环", "相机掉线、图像异常、幅宽异常报警和确认流程", "____", "____", "异常可报警、可记录、可确认或拦截"],
            ["批号与数据安全", "批号唯一性、冲突防护、受控修改、操作审计", "____", "____", "批号冲突不覆盖、不删除历史数据"],
            ["幅宽检测与追溯", "上下限判断、异常报警、持久化、历史查询", "____", "____", "幅宽数据可报警、可查询、可追溯"],
            ["报表与历史查询", "多维查询、异常统计、基础导出、查询性能优化", "____", "____", "统计报表和导出功能可用"],
            ["界面与参数操作", "高频参数入口、参数分组、模板和常用操作简化", "____", "____", "操作界面整改版可现场使用"],
            ["算法与模型部署", "标定板、显卡环境、数据标注、训练、部署和推理验证", "____", "____", "标定、推理、缺陷定位链路可运行"],
            ["测试、验收与文档", "模块测试、全流程联调、操作说明、故障处理说明", "____", "____", "现场验证版、初版文档和问题清单"],
        ],
        [1.25, 2.18, 0.82, 0.92, 2.03],
        font_size=7.45,
        header_size=7.8,
    )

    doc.core_properties.title = "熔喷视觉检测软件重新设计计划（4页版）"
    doc.core_properties.subject = "控制在四页内的软件设计计划与人员分工表"
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    build_doc()
